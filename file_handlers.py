"""
file_handlers.py - معالجة أنواع الملفات المختلفة
"""

import os
from typing import Tuple
from pathlib import Path
from docx import Document
import pdfplumber
from core import TextAnonymizerCore

class FileHandler:
    """
    فئة أساسية لمعالجة الملفات.
    """

    @staticmethod
    def extract_text(file_path: str) -> str:
        """
        استخراج النص من الملف حسب الصيغة.
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"الملف غير موجود: {file_path}")
        
        ext = Path(file_path).suffix.lower()
        try:
            if ext == '.txt':
                with open(file_path, 'r', encoding='utf-8') as f:
                    return f.read()
            elif ext == '.docx':
                doc = Document(file_path)
                paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
                return '\n'.join(paragraphs)
            elif ext == '.pdf':
                text = ""
                with pdfplumber.open(file_path) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + '\n'
                return text.strip()
            else:
                raise ValueError("صيغة غير مدعومة")
        except Exception as e:
            raise RuntimeError(f"خطأ في استخراج النص: {e}")

    @staticmethod
    def process_file(input_path: str, entity_types: List[str], replacement_mode: str, output_dir: str) -> Tuple[str, str]:
        """
        معالجة الملف بالكامل: استخراج -> إخفاء -> حفظ.
        """
        # 1. استخراج النص
        text = FileHandler.extract_text(input_path)
        
        # 2. إخفاء الهوية
        core = TextAnonymizerCore()  # يمكن تحسين الأداء بتمرير كائن Core كمعامل
        anonymized_text = core.anonymize_text(text, entity_types, replacement_mode)
        
        # 3. حفظ الناتج
        os.makedirs(output_dir, exist_ok=True)
        base_name = Path(input_path).stem
        output_txt_path = os.path.join(output_dir, f"{base_name}_anonymized.txt")
        
        with open(output_txt_path, 'w', encoding='utf-8') as f:
            f.write(anonymized_text)
        
        # 4. إعادة بناء .docx إذا لزم الأمر
        if input_path.lower().endswith('.docx'):
            output_docx_path = os.path.join(output_dir, f"{base_name}_anonymized.docx")
            FileHandler.rebuild_docx(input_path, anonymized_text, output_docx_path)
            return output_docx_path, anonymized_text
        
        return output_txt_path, anonymized_text

    @staticmethod
    def rebuild_docx(input_path: str, anonymized_text: str, output_path: str):
        """
        إعادة بناء ملف .docx مع الحفاظ على التنسيق الأساسي.
        """
        try:
            doc = Document(input_path)
            anon_paragraphs = anonymized_text.split('\n')
            
            for i, para in enumerate(doc.paragraphs):
                if i < len(anon_paragraphs):
                    if para.runs:
                        para.runs[0].text = anon_paragraphs[i]
                        for run in para.runs[1:]:
                            run.text = ""
                    else:
                        para.text = anon_paragraphs[i]
                else:
                    para.text = ""
            
            doc.save(output_path)
        except Exception as e:
            raise RuntimeError(f"خطأ في إعادة بناء .docx: {e}")